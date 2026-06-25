from docx import Document
from docx.shared import Pt

doc = Document("C:/Course/Third/SoftwareEngineeringTraining/project_code/实训报告_徐逸_202304020417.docx")

cell = doc.tables[0].rows[0].cells[1]
cell.text = ""

learning = """本次课程设计让我对软件工程全栈开发有了完整而深入的理解与实践。我负责实验室与场所管理模块和设施设备管理模块的开发，从数据库设计、后端CRUD接口开发到前端页面实现，完整经历了一个功能模块从零到一的开发过程。

在数据库设计阶段，我学习了如何根据业务需求设计合理的表结构，包括字段类型选择、主外键约束、索引设计等。lab、lab_booking、equipment、maintenance_record四张表的设计让我对关系型数据库的规范化设计有了更深刻的认识。

在后端开发中，通过使用 Spring Boot 3 + MyBatis-Plus + Sa-Token 技术栈，我掌握了标准化的 CRUD 开发流程，包括 Controller 层的 RESTful 接口设计、Service 层的业务逻辑封装、Mapper 层的数据库操作。同时深入理解了依赖注入、事务管理、统一异常处理等 Spring Boot 核心概念。

在前端开发中，使用 Vue 3 Composition API + Element Plus 组件库，我学会了组件化开发和页面状态管理。通过开发实验室管理、设备管理等前端页面，我熟悉了 el-table、el-dialog、el-form、el-tag 等 Element Plus 组件的使用，以及 Vue Router 路由配置和 Pinia 状态管理。

开发过程中遇到了各种实际问题，如文件编码导致的编译错误、中文乱码问题、图片文件路径配置问题等。通过逐一排查和解决这些问题，我的代码调试能力和问题分析能力得到了显著提升。

本次实训不仅锻炼了我的技术能力，也让我认识到规范编码、版本控制、文档记录在团队协作中的重要性。"""

cell.text = learning + "\n\n\n\n\n学生（签名）：\n（备注：手写签名）"

for p in cell.paragraphs:
    for run in p.runs:
        run.font.size = Pt(11)

doc.save("C:/Course/Third/SoftwareEngineeringTraining/project_code/实训报告_徐逸_202304020417.docx")
print("学习心得已填写完成")
