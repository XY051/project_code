from __future__ import annotations

import base64
import html
import math
import struct
import subprocess
import tempfile
import wave
from pathlib import Path

from docx import Document
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
FILES_DIR = ROOT / "backend" / "files"
BASE_URL = "http://localhost:8080/api/file/files"


RESOURCES = {
    27: {
        "slug": "teaching-video-dunhuang",
        "category_id": 4,
        "title": "敦煌莫高窟艺术导学视频",
        "type": "视频",
        "subtitle": "文化艺术微课：从壁画、造像和丝路交流认识敦煌",
        "cover": "敦煌艺术导学",
        "summary": "围绕莫高窟壁画、佛教造像和丝路文化交流，帮助学生建立文化艺术类资源的学习框架。",
        "intro": "从壁画、造像和丝路交流三个角度认识敦煌文化。",
        "video_text": "同学们好。本视频带大家认识敦煌莫高窟。第一，观察壁画中的线条、色彩和故事场景。第二，理解造像比例、姿态和宗教审美。第三，把敦煌放回丝绸之路的交流背景，思考中西文化如何在图像和工艺中相遇。学习结束后，请选择一幅壁画，记录它的主题、人物关系和色彩特点。",
        "video_slides": [
            ("敦煌莫高窟艺术导学", ["观察壁画构图", "理解造像审美", "连接丝路交流"]),
            ("壁画怎么看", ["先看人物与场景", "再看线条和色彩", "最后看故事主题"]),
            ("造像怎么看", ["比例与姿态", "服饰与手势", "光影与空间感"]),
            ("丝路交流", ["图像题材的传播", "工艺与颜料流动", "多元文化的融合"]),
            ("课堂任务", ["选择一幅壁画做观察记录", "写出主题、色彩和人物关系", "用 3 分钟进行小组分享"]),
        ],
        "sections": [
            ("学习目标", ["能够从图像、造像和历史交流三个角度观察敦煌艺术。", "理解文化遗产数字化资源在课堂导学中的作用。", "完成一份壁画观察记录。"]),
            ("课堂任务", ["播放视频后，学生分组选取一幅壁画进行视觉要素分析。", "教师引导学生把图像细节与丝路交流背景联系起来。", "小组用 3 分钟分享观察结果。"]),
        ],
    },
    28: {
        "slug": "teaching-video-science-communication",
        "category_id": 5,
        "title": "科学传播短视频导学",
        "type": "视频",
        "subtitle": "新闻传播微课：用航天科普案例组织科学信息",
        "cover": "科学传播导学",
        "summary": "以航天科普为例，讲解科学短视频如何完成选题、事实核验、脚本结构和可视化表达。",
        "intro": "以航天科普为例，学习科学信息采集、核验、表达和发布。",
        "video_text": "本视频示范科学传播短视频的制作思路。第一步，确定一个清楚的问题，例如火箭为什么要分级。第二步，查找可靠来源，核对关键数字和术语。第三步，把脚本分为现象、原理、案例和总结四段。第四步，用图示、关键词和旁白降低理解难度。请同学们课后选择一个科学问题，写出一分钟短视频脚本。",
        "video_slides": [
            ("科学传播短视频导学", ["确定科学问题", "核验事实来源", "组织一分钟脚本"]),
            ("选题", ["问题要具体", "观众能理解", "课堂任务可完成"]),
            ("核验", ["记录来源", "核对术语", "避免夸张结论"]),
            ("表达", ["现象引入", "原理解释", "案例支撑", "一句话总结"]),
            ("课堂任务", ["选择一个科学问题", "写一分钟脚本", "标出两个可靠来源"]),
        ],
        "sections": [
            ("学习目标", ["理解科学传播短视频的基本结构。", "能够区分事实、解释和观点。", "完成一个一分钟科普脚本。"]),
            ("课堂任务", ["学生围绕一个科学问题写短视频提纲。", "教师检查事实来源和术语表达。", "小组互评脚本是否清晰、准确、有节奏。"]),
        ],
    },
    29: {
        "slug": "teaching-audio-platform-guide",
        "category_id": 3,
        "title": "虚拟仿真实训平台使用讲解音频",
        "type": "音频",
        "subtitle": "音频导学：平台登录、资源查找与实验提交",
        "cover": "平台使用讲解",
        "summary": "帮助学生快速了解平台登录、资源筛选、实验入口和学习记录提交方式。",
        "intro": "面向学生的平台使用说明音频。",
        "audio_text": "同学们好。本段音频讲解虚拟仿真实训平台的基本使用方法。第一步，登录平台并确认个人信息。第二步，在资源中心按资源类型和专业大类筛选课程资源。第三步，虚拟仿真资源需要先提交申请，审核通过后再进入实验。视频、音频和文档资源可以直接查看。第四步，完成学习后，请保存截图、实验数据和心得，按教师要求提交。",
        "sections": [
            ("收听任务", ["记录平台入口、资源筛选、实验提交三个关键词。", "听完后完成个人学习计划。"]),
            ("课堂建议", ["教师可在课前 5 分钟播放本音频。", "学生边听边检查自己的账号和浏览器环境。"]),
        ],
    },
    30: {
        "slug": "teaching-audio-open-apply",
        "category_id": 5,
        "title": "共享开放申请流程讲解音频",
        "type": "音频",
        "subtitle": "音频导学：共享开放资源申请与审核流程",
        "cover": "申请流程讲解",
        "summary": "说明资源申请、管理员审核、资源占用、使用完成和归还的完整流程。",
        "intro": "讲解资源申请、审核和使用流程。",
        "audio_text": "本音频介绍共享开放申请流程。第一，选择需要开放使用的虚拟仿真资源，查看可用数量和开放类型。第二，填写使用目的、计划时间和联系方式，提交申请。第三，管理员审核通过后，系统占用一个资源名额，学生可以进入实验。第四，使用完成后进行归还或由管理员确认，资源数量恢复。请注意，资源数量为零时，需要等待管理员协调。",
        "sections": [
            ("收听任务", ["说出申请、审核、使用、归还四个流程节点。", "理解资源数量为 0 时为什么不能立即进入实验。"]),
            ("课堂建议", ["适合共享开放说明会或课前申请指导。", "可配合申请页面演示一起使用。"]),
        ],
    },
    31: {
        "slug": "teaching-doc-training-guide",
        "category_id": 1,
        "title": "虚拟仿真实训指导书",
        "type": "文档",
        "subtitle": "教学文档：任务书、实验步骤与评价标准",
        "cover": "实训指导书",
        "summary": "提供虚拟仿真实训的任务说明、实验步骤、数据记录表和评价标准。",
        "intro": "实训任务、实验步骤和考核要求文档。",
        "sections": [
            ("一、实训目标", ["理解虚拟仿真实训的业务背景和任务要求。", "能够独立完成平台操作、数据记录和结果分析。", "形成规范化实训报告。"]),
            ("二、实验步骤", ["登录平台，进入资源中心。", "选择对应资源并阅读资源说明。", "虚拟仿真资源先提交申请，审核后进入实验。", "记录关键参数、操作截图和异常问题。", "完成心得总结并提交给教师。"]),
            ("三、评价标准", ["过程记录完整，占 30%。", "操作结果正确，占 40%。", "问题分析和改进建议，占 20%。", "团队协作与课堂表现，占 10%。"]),
        ],
    },
    32: {
        "slug": "teaching-doc-open-apply",
        "category_id": 4,
        "title": "实训资源开放申请说明",
        "type": "文档",
        "subtitle": "教学文档：共享开放申请条件、流程和注意事项",
        "cover": "开放申请说明",
        "summary": "说明共享开放资源的申请条件、审核材料、使用规则和常见问题。",
        "intro": "共享开放申请条件、流程和注意事项。",
        "sections": [
            ("一、申请条件", ["申请人需要具备平台账号。", "申请用途应与课程实训、项目训练或开放学习相关。", "虚拟仿真资源需有可用名额。"]),
            ("二、申请流程", ["选择资源并进入详情页。", "点击申请资源并填写使用目的。", "等待管理员审核。", "审核通过后按要求进入实验并完成记录。"]),
            ("三、注意事项", ["视频、音频、文档资源可直接查看，不需要提交资源申请。", "虚拟仿真资源使用完成后应按要求归还。", "资源数量不足时请联系管理员协调。"]),
        ],
    },
}


CSS = """
body{margin:0;background:#f4f7fb;color:#172033;font-family:"Microsoft YaHei",Arial,sans-serif;}
.top{height:32px;background:linear-gradient(100deg,#123075,#1b70d8 46%,#47b0e8 62%,#5b2ab5);}
.wrap{max-width:1040px;margin:0 auto;padding:34px 22px 56px;}
.hero{background:#fff;border:1px solid #e6edf5;border-radius:10px;padding:30px;box-shadow:0 10px 24px rgba(21,47,74,.05);}
.badge{display:inline-flex;padding:6px 12px;border-radius:999px;background:#e8f2ff;color:#1478d4;font-weight:700;font-size:14px;}
h1{margin:18px 0 10px;font-size:32px;line-height:1.25;}
.subtitle{margin:0;color:#526173;font-size:17px;line-height:1.8;}
.media{margin-top:24px;background:#f8fafc;border:1px solid #edf2f7;border-radius:10px;padding:18px;}
.poster{display:grid;place-items:center;min-height:230px;border-radius:8px;background:linear-gradient(135deg,#17324d,#2563eb 55%,#2f7d68);color:white;font-size:28px;font-weight:800;}
.video{display:block;width:100%;aspect-ratio:16/9;border-radius:8px;background:#0f172a;}
.audio{width:100%;margin-top:14px;}
.media-note{margin:12px 0 0;color:#64748b;font-size:14px;}
.section{margin-top:20px;background:#fff;border:1px solid #e6edf5;border-radius:10px;padding:24px;}
h2{margin:0 0 12px;color:#1478d4;font-size:22px;}
ul{margin:0;padding-left:22px;}li{margin:8px 0;line-height:1.75;}
.summary{font-size:16px;line-height:1.9;color:#303846;}
.script{white-space:pre-wrap;line-height:1.9;color:#303846;background:#f8fafc;border:1px solid #edf2f7;border-radius:8px;padding:14px;}
"""


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size=size)
    return ImageFont.load_default()


def draw_wrapped_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    x: int,
    y: int,
    max_width: int,
    font_obj: ImageFont.FreeTypeFont,
    fill: str,
    line_gap: int = 10,
) -> int:
    line = ""
    for char in text:
        candidate = line + char
        if draw.textlength(candidate, font=font_obj) <= max_width:
            line = candidate
        else:
            draw.text((x, y), line, fill=fill, font=font_obj)
            y += font_obj.size + line_gap
            line = char
    if line:
        draw.text((x, y), line, fill=fill, font=font_obj)
        y += font_obj.size + line_gap
    return y


def make_cover(slug: str, title: str, label: str, kind: str) -> str:
    path = FILES_DIR / f"{slug}.png"
    img = Image.new("RGB", (1280, 720), "#17324d")
    draw = ImageDraw.Draw(img)
    for y in range(720):
        r = int(18 + y * 0.03)
        g = int(48 + y * 0.10)
        b = int(117 + y * 0.12)
        draw.line((0, y, 1280, y), fill=(r, min(g, 140), min(b, 220)))
    draw.rounded_rectangle((70, 70, 1210, 650), radius=32, outline=(255, 255, 255), width=3)
    draw.text((105, 120), kind, fill="#c8e7ff", font=font(34, True))
    draw_wrapped_text(draw, label, 105, 238, 1040, font(62, True), "white", 12)
    draw_wrapped_text(draw, title, 105, 360, 1040, font(42, True), "#eaf6ff", 12)
    draw.text((105, 560), "共享开放教学资源", fill="#b8d7ef", font=font(28))
    img.save(path)
    return path.name


def make_wav(slug: str, text: str) -> str:
    wav_path = FILES_DIR / f"{slug}.wav"
    command = f"""
$voice = New-Object -ComObject SAPI.SpVoice
$stream = New-Object -ComObject SAPI.SpFileStream
$stream.Open('{str(wav_path).replace("'", "''")}', 3)
$voice.AudioOutputStream = $stream
$voice.Rate = 0
$voice.Volume = 100
$voice.Speak(@'
{text}
'@) | Out-Null
$stream.Close()
"""
    encoded = base64.b64encode(command.encode("utf-16le")).decode("ascii")
    subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-EncodedCommand", encoded],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        check=False,
    )
    if not wav_path.exists() or wav_path.stat().st_size < 1000:
        framerate = 22050
        seconds = 8
        with wave.open(str(wav_path), "w") as wav:
            wav.setnchannels(1)
            wav.setsampwidth(2)
            wav.setframerate(framerate)
            for i in range(framerate * seconds):
                value = int(9000 * math.sin(2 * math.pi * 440 * i / framerate))
                wav.writeframes(struct.pack("<h", value))
    return wav_path.name


def ffmpeg_exe() -> str:
    try:
        import imageio_ffmpeg

        return imageio_ffmpeg.get_ffmpeg_exe()
    except Exception:
        return "ffmpeg"


def make_video_slide(path: Path, item: dict, index: int, title: str, bullets: list[str]) -> None:
    palettes = [
        ("#102a43", "#2563eb", "#2dd4bf"),
        ("#301934", "#7c3aed", "#f59e0b"),
        ("#17324d", "#0ea5e9", "#22c55e"),
        ("#1f2937", "#dc2626", "#facc15"),
        ("#0f172a", "#14b8a6", "#a78bfa"),
    ]
    base, accent, glow = palettes[index % len(palettes)]
    img = Image.new("RGB", (1280, 720), base)
    draw = ImageDraw.Draw(img)
    for y in range(720):
        ratio = y / 720
        r1, g1, b1 = tuple(int(base[i : i + 2], 16) for i in (1, 3, 5))
        r2, g2, b2 = tuple(int(accent[i : i + 2], 16) for i in (1, 3, 5))
        fill = (
            int(r1 * (1 - ratio) + r2 * ratio * 0.75),
            int(g1 * (1 - ratio) + g2 * ratio * 0.75),
            int(b1 * (1 - ratio) + b2 * ratio * 0.75),
        )
        draw.line((0, y, 1280, y), fill=fill)

    draw.rounded_rectangle((72, 70, 1208, 650), radius=28, outline=(255, 255, 255), width=3)
    draw.ellipse((930, 95, 1160, 325), outline=glow, width=8)
    draw.ellipse((980, 145, 1110, 275), outline=(255, 255, 255), width=3)
    draw.line((95, 560, 1180, 560), fill=(255, 255, 255), width=2)
    draw.text((105, 110), item["type"] + "资源", fill="#dbeafe", font=font(32, True))
    y = draw_wrapped_text(draw, title, 105, 190, 760, font(60, True), "white", 14)
    y += 24
    for bullet in bullets:
        draw.rounded_rectangle((112, y + 10, 136, y + 34), radius=12, fill=glow)
        draw_wrapped_text(draw, bullet, 156, y, 760, font(34, True), "#f8fafc", 10)
        y += 62
    draw.text((105, 590), "资源中心 · 可直接播放 · 无需申请", fill="#dbeafe", font=font(26))
    img.save(path)


def make_video(item: dict, cover_name: str) -> str:
    video_name = f'{item["slug"]}.mp4'
    video_path = FILES_DIR / video_name
    audio_name = make_wav(f'{item["slug"]}-narration', item["video_text"])
    audio_path = FILES_DIR / audio_name
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        concat_path = temp_path / "slides.txt"
        slide_paths = []
        for index, (title, bullets) in enumerate(item["video_slides"]):
            slide_path = temp_path / f"slide-{index}.png"
            make_video_slide(slide_path, item, index, title, bullets)
            slide_paths.append(slide_path)

        lines = []
        for slide_path in slide_paths:
            normalized = slide_path.as_posix()
            lines.append(f"file '{normalized}'")
            lines.append("duration 7")
        lines.append(f"file '{slide_paths[-1].as_posix()}'")
        concat_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

        subprocess.run(
            [
                ffmpeg_exe(),
                "-y",
                "-f",
                "concat",
                "-safe",
                "0",
                "-i",
                str(concat_path),
                "-i",
                str(audio_path),
                "-vf",
                "fps=24,format=yuv420p",
                "-c:v",
                "libx264",
                "-preset",
                "veryfast",
                "-c:a",
                "aac",
                "-b:a",
                "128k",
                "-shortest",
                "-movflags",
                "+faststart",
                str(video_path),
            ],
            check=True,
        )
    return video_name


def render_sections(sections: list[tuple[str, list[str]]]) -> str:
    blocks = []
    for title, items in sections:
        lines = "".join(f"<li>{html.escape(item)}</li>" for item in items)
        blocks.append(f'<section class="section"><h2>{html.escape(title)}</h2><ul>{lines}</ul></section>')
    return "\n".join(blocks)


def make_html(
    item: dict,
    cover_name: str,
    video_name: str | None = None,
    audio_name: str | None = None,
) -> str:
    slug = item["slug"]
    path = FILES_DIR / f"{slug}.html"
    media = f'<div class="poster">{html.escape(item["cover"])}</div>'
    if video_name:
        media = (
            f'<video class="video" controls preload="metadata" poster="{cover_name}" src="{video_name}"></video>'
            f'<p class="media-note">如果浏览器没有自动加载，请点击播放器左下角播放按钮。</p>'
        )
    elif audio_name:
        media += f'<audio class="audio" controls src="{audio_name}"></audio>'
        media += f'<div class="script">{html.escape(item["audio_text"])}</div>'
    body = f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{html.escape(item["title"])}</title>
<style>{CSS}</style>
</head>
<body>
<div class="top"></div>
<main class="wrap">
  <section class="hero">
    <span class="badge">{html.escape(item["type"])}资源</span>
    <h1>{html.escape(item["title"])}</h1>
    <p class="subtitle">{html.escape(item["subtitle"])}</p>
    <div class="media">{media}</div>
  </section>
  <section class="section">
    <h2>资源概述</h2>
    <p class="summary">{html.escape(item["summary"])}</p>
  </section>
  {render_sections(item["sections"])}
</main>
</body>
</html>"""
    path.write_text(body, encoding="utf-8")
    return path.name


def make_docx(item: dict) -> str:
    path = FILES_DIR / f'{item["slug"]}.docx'
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"].font.size = Pt(11)
    doc.add_heading(item["title"], level=0)
    doc.add_paragraph(item["subtitle"])
    doc.add_heading("资源概述", level=1)
    doc.add_paragraph(item["summary"])
    for title, lines in item["sections"]:
        doc.add_heading(title, level=1)
        for line in lines:
            doc.add_paragraph(line, style="List Bullet")
    doc.save(path)
    return path.name


def sql_value(value: str) -> str:
    return "'" + value.replace("\\", "\\\\").replace("'", "''") + "'"


def main() -> None:
    FILES_DIR.mkdir(parents=True, exist_ok=True)
    updates = []
    for resource_id, item in RESOURCES.items():
        cover_name = make_cover(item["slug"], item["title"], item["cover"], item["type"])
        video_name = make_video(item, cover_name) if item["type"] == "视频" else None
        audio_name = make_wav(item["slug"], item["audio_text"]) if item["type"] == "音频" else None
        html_name = make_html(item, cover_name, video_name, audio_name)
        docx_name = make_docx(item) if item["type"] == "文档" else None

        description = item["summary"]
        if video_name:
            description += f'<p><a href="{BASE_URL}/{video_name}" target="_blank">直接打开 MP4 视频</a></p>'
        if audio_name:
            description += f'<p><a href="{BASE_URL}/{audio_name}" target="_blank">直接收听音频</a></p>'
        if docx_name:
            description += f'<p><a href="{BASE_URL}/{docx_name}" target="_blank">下载 Word 教学文档</a></p>'

        updates.append(
            "UPDATE training_resource SET "
            f"name={sql_value(item['title'])}, "
            f"category_id={item['category_id']}, "
            f"resource_type={sql_value(item['type'])}, "
            "open_type='开放共享', "
            "resource_level='在线资源', "
            f"cover_image={sql_value('/file/files/' + cover_name)}, "
            f"intro={sql_value(item['intro'])}, "
            f"access_url={sql_value(BASE_URL + '/' + html_name)}, "
            f"description={sql_value(description)}, "
            "available_count=0, status='可查看' "
            f"WHERE id={resource_id};"
        )
    sql_path = ROOT / "backend" / "sql" / "update_teaching_resources.sql"
    sql_path.write_text("\n".join(updates) + "\n", encoding="utf-8")
    print(sql_path)


if __name__ == "__main__":
    main()
